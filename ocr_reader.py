# ocr_reader.py
import sys
import os
import argparse
from PIL import Image
import pytesseract
from docx import Document

def extract_text(image_path, lang='eng'):
    """Extract text from image using Tesseract OCR."""
    try:
        img = Image.open(image_path)
        text = pytesseract.image_to_string(img, lang=lang)
        return text.strip()
    except Exception as e:
        print(f"Error processing {image_path}: {e}", file=sys.stderr)
        return None

def export_docx(text, output_path):
    """Export extracted text to DOCX format."""
    doc = Document()
    doc.add_heading('OCR Extracted Text', 0)
    doc.add_paragraph(text)
    doc.save(output_path)
    return output_path

def export_txt(text, output_path):
    """Export extracted text to plain text format."""
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(text)
    return output_path

def main():
    parser = argparse.ArgumentParser(description="OCR Reader")
    parser.add_argument('images', nargs='+', help='Image files to process')
    parser.add_argument('--output', default='output.docx', help='Output file')
    parser.add_argument('--lang', default='eng', help='OCR language (default: eng)')
    parser.add_argument('--format', choices=['docx', 'txt'], default='docx', help='Output format')
    args = parser.parse_args()

    print("📄 OCR Reader")
    print(f"Images: {', '.join(args.images)}")
    print(f"Language: {args.lang}")

    all_text = []
    for img_path in args.images:
        print(f"Processing: {img_path}")
        text = extract_text(img_path, args.lang)
        if text:
            all_text.append(f"--- {img_path} ---\n{text}\n")
        else:
            print(f"⚠️ No text extracted from {img_path}")

    if not all_text:
        print("❌ No text extracted from any image.")
        sys.exit(1)

    combined_text = "\n".join(all_text)
    print("\n📖 Extracted text:")
    print(combined_text[:500] + ("..." if len(combined_text) > 500 else ""))

    if args.format == 'docx':
        output_path = export_docx(combined_text, args.output)
        print(f"\n📄 Exporting to DOCX... ✅")
    else:
        output_path = export_txt(combined_text, args.output)
        print(f"\n📄 Exporting to TXT... ✅")
    print(f"File saved: {output_path}")

if __name__ == "__main__":
    main()
